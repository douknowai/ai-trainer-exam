#!/usr/bin/env python3
"""生成「人工智能训练师五级练习与考试系统 - 使用说明书」Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "docs",
                      "人工智能训练师五级练习与考试系统-使用说明书.docx")

# ── helpers ──────────────────────────────────────────────────

PRIMARY = RGBColor(0x1a, 0x5c, 0x4b)   # 墨青绿
ACCENT  = RGBColor(0xb0, 0x6e, 0x2a)   # 暖橙
DARK    = RGBColor(0x33, 0x33, 0x33)
GRAY    = RGBColor(0x66, 0x66, 0x66)


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex
    })
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY if level <= 2 else DARK
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h


def add_body(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1A5C4B')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if col_widths:
        for ri in range(len(rows) + 1):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)
    return table


# ── build document ───────────────────────────────────────────

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# default font
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── cover ────────────────────────────────────────────────────

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('人工智能训练师五级\n练习与考试系统')
run.font.size = Pt(32)
run.font.color.rgb = PRIMARY
run.bold = True
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('使 用 说 明 书')
run.font.size = Pt(20)
run.font.color.rgb = ACCENT
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('版本 2.0\n2026 年 7 月')
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ── TOC placeholder ──────────────────────────────────────────

add_heading(doc, '目  录', level=1)
add_body(doc, '（请在 Word 中右键此处 → 更新域 → 更新整个目录）')
# real TOC field
paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar1 = doc.element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
run._element.append(fldChar1)
instrText = doc.element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
run._element.append(instrText)
fldChar2 = doc.element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
run._element.append(fldChar2)
fldChar3 = doc.element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
run._element.append(fldChar3)

doc.add_page_break()

print("Skeleton built, appending sections...")

# ═════════════════════════════════════════════════════════════
# 第一章  系统概述
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第一章  系统概述', level=1)

add_heading(doc, '1.1  产品简介', level=2)
add_body(doc,
    '本系统是面向「人工智能训练师五级」职业资格培训的零基础练习与考试平台。'
    '系统覆盖理论练习、实操任务训练、正式考试、成绩管理与复核全流程，'
    '服务于三类用户角色：零基础学员、教师、管理管理员。')

add_heading(doc, '1.2  适用对象', level=2)
add_table(doc,
    ['角色', '适用人群', '核心职责'],
    [
        ['学员', '参加五级考试的零基础培训学员', '刷题练习、实操训练、参加考试、查看成绩'],
        ['教师', '授课教师、班主任', '查看学员进度、管理考试、查看成绩'],
        ['管理员', '学校管理员、超级管理员', '管理组织/班级/用户/题库/考务/成绩/系统配置'],
    ],
    col_widths=[2.5, 5, 8])

add_heading(doc, '1.3  运行环境', level=2)
add_body(doc, '推荐在以下环境中使用：')
add_table(doc,
    ['项目', '要求'],
    [
        ['操作系统', 'Windows 10 及以上 / macOS 12 及以上'],
        ['浏览器', 'Chrome 110+ 或 Edge 110+（不建议使用 IE）'],
        ['屏幕分辨率', '1366×768 及以上（推荐 1920×1080）'],
        ['网络', '稳定的互联网连接'],
        ['外设', '耳机/麦克风（音频转写任务需要）'],
    ],
    col_widths=[4, 11])

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 第二章  账号与登录
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第二章  账号与登录', level=1)

add_heading(doc, '2.1  登录方式', level=2)
add_body(doc, '系统使用邮箱 + 密码方式登录。打开浏览器，在地址栏输入系统网址，进入登录页面。')
add_body(doc, '操作步骤：', bold=True)
add_bullet(doc, '输入注册邮箱地址')
add_bullet(doc, '输入密码')
add_bullet(doc, '点击「登录」按钮')
add_bullet(doc, '系统根据您的角色自动跳转到对应首页（学员/教师/管理员）')

add_heading(doc, '2.2  测试账号', level=2)
add_body(doc, '系统预置了以下测试账号供培训与验收使用：')
add_body(doc, '以下为开发测试阶段的账号密码，正式上线前请务必修改：', bold=True, color=ACCENT)
add_table(doc,
    ['角色', '邮箱', '密码', '权限说明'],
    [
        ['超级管理员', 'admin@exam.local', 'Admin@2026', '拥有全部权限，可管理所有模块'],
        ['学校管理员', 'school@exam.local', 'School@2026', '管理本校事务（用户/题库/考务/成绩）'],
        ['教师', 'teacher01@exam.local', 'Teacher@2026', '查看所辖班级学员进度与成绩'],
        ['监考员', 'invig01@exam.local', 'Invig@2026', '考试监控与成绩查看'],
        ['审计员', 'auditor01@exam.local', 'Audit@2026', '查看审计日志与报表'],
        ['题库编辑', 'editor01@exam.local', 'Editor@2026', '编辑练习/考试题库'],
        ['题库审核', 'reviewer01@exam.local', 'Review@2026', '审核待发布题目'],
        ['学员 001', 'stu001@student.exam.local', 'abcd2345', '练习与考试（主测试账号）'],
        ['学员 002', 'stu002@student.exam.local', 'efgh6789', '练习与考试（辅助测试账号）'],
    ],
    col_widths=[2.8, 5.5, 3.2, 4])

add_heading(doc, '2.3  密码安全', level=2)
add_bullet(doc, '首次部署后请立即修改管理员密码')
add_bullet(doc, '学员密码由教师/管理员统一分配，可批量重置')
add_bullet(doc, '密码以加密形式存储，管理员无法查看学员原始密码')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 第三章  学员端使用指南
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第三章  学员端使用指南', level=1)

add_heading(doc, '3.1  学员首页', level=2)
add_body(doc, '路径：/student/home', bold=True, color=PRIMARY)
add_body(doc,
    '登录后进入学员首页，展示学习概览信息：')
add_bullet(doc, '练习进度统计（已练习题数、正确率）')
add_bullet(doc, '实操任务完成数')
add_bullet(doc, '可参加的考试（如有）')
add_bullet(doc, '最近成绩')
add_body(doc,
    '顶部导航栏提供以下入口：首页、理论练习、实操练习、错题本、考试、成绩、帮助。',
    bold=True)

add_heading(doc, '3.2  理论练习', level=2)
add_body(doc, '路径：/student/practice', bold=True, color=PRIMARY)
add_body(doc, '理论练习采用「驾考科目一」模式——刷题备考是设计目标。')
add_body(doc, '功能说明：', bold=True)
add_bullet(doc, '逐题练习：显示题目和选项，选择后立即给出对错反馈')
add_bullet(doc, '答对显示绿色「✓ 做对了！」，答错显示红色「✗ 答错了，没关系」并展示正确答案')
add_bullet(doc, '练习中可以看到正确答案，这是设计行为（非缺陷），帮助学员学习')
add_bullet(doc, '做错的题目自动加入错题本')

add_heading(doc, '3.3  实操练习', level=2)
add_body(doc, '路径：/student/task', bold=True, color=PRIMARY)
add_body(doc, '实操任务模拟真实数据标注工作场景，涵盖以下类型：')

add_table(doc,
    ['任务类型', '操作说明', '评分方式'],
    [
        ['Excel 数据清洗', '识别并删除不符合要求的数据行', '按正确删除/误删行数计分'],
        ['统计表格', '在虚拟表格中完成指定操作', '按操作正确率计分'],
        ['文件分类', '将文件拖放到正确的分类文件夹', '按放对文件数计分'],
        ['图片清洗', '从一组图片中剔除不合格样本（如足球混入篮球集）', '按正确删除/误删数计分'],
        ['图片标注-方框', '在图片上用矩形框标出指定目标', '按 IoU（交并比）计分'],
        ['图片标注-点', '在图片上点击指定目标位置', '按点击距离精度计分'],
        ['图片标注-折线', '沿目标边缘画折线', '按 Chamfer 距离计分'],
        ['图片标注-多边形', '用多边形圈出目标区域', '按多边形 IoU 计分'],
        ['文本情感标注', '判断文本是正向/中性/负面', '按标注准确率计分'],
        ['音频转写', '听音频并输入文字内容', '按字符错误率(CER)计分'],
        ['数据标注综合', '组合多项操作的综合任务', '按子任务权重加权计分'],
    ],
    col_widths=[3.5, 7, 4.5])

add_heading(doc, '3.4  错题本', level=2)
add_body(doc, '路径：/student/wrong', bold=True, color=PRIMARY)
add_body(doc,
    '自动收集理论练习中答错的题目，方便针对性复习。'
    '可以重新做错题，做对后自动移出错题本。')

add_heading(doc, '3.5  参加考试', level=2)
add_body(doc, '路径：/student/exams', bold=True, color=PRIMARY)
add_body(doc, '考试功能说明：', bold=True)
add_bullet(doc, '考试列表展示所有已安排且在有效期内的考试')
add_bullet(doc, '点击「开始考试」后，服务端校验考试时间，创建考试记录')
add_bullet(doc, '考试界面顶部显示倒计时，时间到自动交卷')
add_bullet(doc, '考试中答案自动保存，刷新页面不会丢失进度，倒计时不重置')
add_bullet(doc, '考试模式界面克制严肃，不显示答题反馈（与练习不同）')
add_body(doc, '注意事项：', bold=True, color=ACCENT)
add_bullet(doc, '考试一旦开始，需在规定时间内完成')
add_bullet(doc, '服务端时间锁保护：超时交卷自动判为「已过期」')
add_bullet(doc, '不可中途暂停或重新开始')

add_heading(doc, '3.6  成绩查询', level=2)
add_body(doc, '路径：/student/results', bold=True, color=PRIMARY)
add_body(doc,
    '查看历史考试的成绩，包括总分、是否通过、各模块得分。'
    '成绩需经管理员发布后才可查看。')

add_heading(doc, '3.7  帮助页面', level=2)
add_body(doc, '路径：/student/help', bold=True, color=PRIMARY)
add_body(doc, '提供系统使用常见问题的帮助说明。')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 第四章  教师端使用指南
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第四章  教师端使用指南', level=1)

add_heading(doc, '4.1  教师仪表盘', level=2)
add_body(doc, '教师登录后查看所管辖班级的概况，包括学员人数、练习进度、考试通过率等。')

add_heading(doc, '4.2  考试管理', level=2)
add_body(doc, '教师可以创建考试并关联到所辖班级，查看考试状态和学员参与情况。')

add_heading(doc, '4.3  学员管理', level=2)
add_body(doc, '查看所辖班级的学员列表、练习进度和成绩。'
    '（注：教师的权限范围限于其被授权的班级，不能查看其他班级数据）')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 第五章  管理端使用指南
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第五章  管理端使用指南', level=1)

add_body(doc, '管理员端采用左侧导航 + 右侧内容区的布局。根据角色权限不同，可见的菜单项也不同。')

add_heading(doc, '5.1  管理端功能总览', level=2)
add_table(doc,
    ['功能模块', '路径', '可用角色', '说明'],
    [
        ['工作台', '/admin/dashboard', '全部', '系统概况与快捷入口'],
        ['学校管理', '/admin/organizations', '超级管理员', '管理机构（学校）'],
        ['培训项目', '/admin/projects', '超管/校管', '管理培训项目'],
        ['班级管理', '/admin/cohorts', '超管/校管', '管理班级、分配教师'],
        ['账号管理', '/admin/users', '超管/校管', '管理用户账号'],
        ['练习题库', '/admin/practice-bank', '超管/校管/编辑/审核', '管理练习题目'],
        ['考试题库', '/admin/exam-bank', '超管/校管/编辑/审核', '管理正式考试题目'],
        ['题库导入', '/admin/import', '超管/校管/编辑', '通过 Word 文件批量导入题目'],
        ['题目审核', '/admin/review', '超管/校管/审核', '审核待发布的题目'],
        ['素材工坊', '/admin/media-studio', '超管/校管/编辑', 'AI 生成图片/音频素材'],
        ['评分校准', '/admin/grading-calibration', '超管/校管', '校准评分参数'],
        ['试卷管理', '/admin/papers', '超管/校管', '组建和管理试卷'],
        ['考务安排', '/admin/exam-schedules', '超管/校管', '安排考试时间/班级/试卷'],
        ['考试监控', '/admin/exam-monitor', '超管/校管/监考', '实时监控考试进度'],
        ['成绩管理', '/admin/results', '超管/校管/监考', '成绩列表/复核/调整/发布'],
        ['审计日志', '/admin/audit', '超管/审计', '系统操作审计记录'],
        ['报表导出', '/admin/reports', '超管/校管/审计', '成绩分布/通过率/班级对比报表'],
        ['系统设置', '/admin/settings', '超级管理员', '系统参数配置'],
    ],
    col_widths=[2.8, 4.2, 4, 4.5])

add_heading(doc, '5.2  成绩复核与调整（重要）', level=2)
add_body(doc, '路径：/admin/results', bold=True, color=PRIMARY)
add_body(doc, '成绩复核是管理员的核心功能之一，包含三个能力：')

add_body(doc, '（1）查看成绩详情', bold=True)
add_bullet(doc, '展示总分 + 6 项分项分数（理论/清洗/标注/文本/音频/统计）')
add_bullet(doc, '逐题对比：学员答案 vs 标准答案 vs 得分 vs 评分详情')

add_body(doc, '（2）复核确认', bold=True)
add_bullet(doc, '复核员确认成绩无误后，填写复核说明')
add_bullet(doc, '点击「确认并发布」将成绩状态改为「已复核」')
add_bullet(doc, '复核操作写入审计日志')

add_body(doc, '（3）手动调整分数', bold=True, color=ACCENT)
add_bullet(doc, '点击「调整分数」按钮进入内联编辑模式')
add_bullet(doc, '可逐项调整 6 个模块的分数，系统自动重算总分')
add_bullet(doc, '调整后的分数实时显示通过/不通过预判')
add_bullet(doc, '必须填写调整原因（至少 5 字），保存后写入审计日志')
add_bullet(doc, '原始分数和调整人信息被保留，可追溯')
add_bullet(doc, '权限限制：仅超级管理员和学校管理员可调整')

add_heading(doc, '5.3  评分系统说明', level=2)
add_body(doc, '系统采用确定性自动评分引擎，提交后即时出分，无需教师人工打分。')

add_table(doc,
    ['评分维度', '算法', '说明'],
    [
        ['理论题（单选/判断）', '精确匹配', '答对满分，答错 0 分'],
        ['矩形框标注', 'IoU 交并比', '阈值 0.45，连续计分'],
        ['点标注', '欧几里得距离', '容差 0.08，越准分越高'],
        ['折线标注', 'Chamfer 距离', '64 点采样对称距离'],
        ['多边形标注', '网格化 IoU', '80×80 网格采样'],
        ['音频转写', '字符错误率 CER', '准确率×0.85 + 语气词召回×0.15'],
        ['数据集质检', 'F1 Score', '精确率与召回率的调和平均'],
    ],
    col_widths=[4, 4, 7])

add_body(doc, '核心设计原则：', bold=True)
add_bullet(doc, '纯函数评分，不依赖网络/LLM/随机数，结果可复现')
add_bullet(doc, '标准答案由服务端冻结，客户端无法篡改')
add_bullet(doc, '最优一对一匹配算法，标注顺序不影响评分结果')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 第六章  常见问题
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第六章  常见问题（FAQ）', level=1)

faqs = [
    ('Q: 考试中途不小心关闭了浏览器怎么办？',
     'A: 重新打开浏览器登录系统，进入考试页面即可恢复。系统自动保存了您的答题进度，'
     '倒计时不会重置，继续计时。'),
    ('Q: 练习时能看到答案，这正常吗？',
     'A: 这是系统设计行为。理论练习采用「驾考刷题」模式，显示答案帮助学员学习备考。'
     '正式考试不会显示答案。'),
    ('Q: 考试时间到了还没做完怎么办？',
     'A: 系统会自动交卷。服务端时间锁会检测超时，超时交卷判为「已过期」，'
     '已完成的答案正常评分。'),
    ('Q: 成绩显示「待发布」是什么意思？',
     'A: 考试结束后成绩需要管理员复核并发布后，学员才能查看。'
     '如果长时间未发布，请联系教师或管理员。'),
    ('Q: 对成绩有异议怎么办？',
     'A: 可以联系教师或管理员，管理员可以在成绩复核页面查看逐题评分详情，'
     '必要时可以手动调整分数。'),
    ('Q: 标注任务怎么算分？是不是要求标得非常精确？',
     'A: 标注任务采用连续计分，不是非对即错。标注越接近标准答案分越高。'
     '系统使用专业算法（IoU/距离/Chamfer）评估精度，有一定容差范围。'),
    ('Q: 忘记密码怎么办？',
     'A: 学员密码请联系教师或管理员重置；管理员密码请联系超级管理员。'),
    ('Q: 标注时拖动不准怎么办？',
     'A: 标注完成后可以拖动调整框/点/线的位置，也可以删除重新标注。'
     '建议使用鼠标操作，精度更高。'),
]

for q, a in faqs:
    add_body(doc, q, bold=True)
    add_body(doc, a)
    doc.add_paragraph()

doc.add_page_break()

# ── save ─────────────────────────────────────────────────────

doc.save(OUTPUT)
print(f"Manual saved to: {OUTPUT}")
print(f"File size: {os.path.getsize(OUTPUT) / 1024:.0f} KB")
