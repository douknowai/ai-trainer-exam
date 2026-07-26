#!/usr/bin/env python3
"""生成「测试评测指导文档」Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "docs",
                      "测试评测指导文档.docx")

PRIMARY = RGBColor(0x1a, 0x5c, 0x4b)
ACCENT  = RGBColor(0xb0, 0x6e, 0x2a)
DARK    = RGBColor(0x33, 0x33, 0x33)
GRAY    = RGBColor(0x66, 0x66, 0x66)
RED     = RGBColor(0xcc, 0x33, 0x33)


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex
    })
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY if level <= 2 else DARK
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h


def add_body(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1A5C4B')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if col_widths:
        for ri in range(len(rows) + 1):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)
    return table


doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── cover ────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('人工智能训练师五级\n练习与考试系统')
run.font.size = Pt(32)
run.font.color.rgb = PRIMARY
run.bold = True
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('测 试 评 测 指 导 文 档')
run.font.size = Pt(20)
run.font.color.rgb = ACCENT
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('供测试人员使用\n2026 年 7 月')
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()
print("Skeleton done")

# ═════════════════════════════════════════════════════════════
# 第一章  测试目的与范围
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第一章  测试目的与范围', level=1)

add_heading(doc, '1.1  测试目的', level=2)
add_body(doc,
    '本文档用于指导测试人员对「人工智能训练师五级练习与考试系统」进行系统性功能测试，'
    '确保各角色（学员、教师、管理员）的核心功能链路正常，发现并记录缺陷，'
    '为系统上线提供质量保证。')

add_heading(doc, '1.2  测试范围', level=2)
add_table(doc,
    ['测试类别', '覆盖内容', '优先级'],
    [
        ['登录认证', '各角色登录、登出、权限校验', 'P0'],
        ['学员端-理论练习', '题目展示、答题判分、错题收集', 'P0'],
        ['学员端-实操任务', '6类实操任务操作与自动评分', 'P0'],
        ['学员端-考试流程', '考试开始→答题→交卷→成绩', 'P0'],
        ['学员端-成绩查询', '成绩列表、详情展示', 'P1'],
        ['教师端', '仪表盘、考试管理、学员查看', 'P1'],
        ['管理端-题库管理', '题目增删改查、导入导出', 'P1'],
        ['管理端-考务管理', '考试安排、试卷管理', 'P1'],
        ['管理端-成绩复核', '查看详情、复核确认、手动调分', 'P0'],
        ['管理端-审计日志', '操作日志记录与查询', 'P2'],
        ['管理端-报表导出', '成绩分布、通过率报表', 'P2'],
        ['管理端-系统设置', '参数配置', 'P2'],
    ],
    col_widths=[4, 8, 2])

add_heading(doc, '1.3  不在本次测试范围', level=2)
add_bullet(doc, '性能/压力测试（后续单独执行）')
add_bullet(doc, '安全渗透测试（需授权环境）')
add_bullet(doc, '移动端兼容性（系统以 PC 1366x768 为主）')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 第二章  测试环境与数据准备
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第二章  测试环境与数据准备', level=1)

add_heading(doc, '2.1  测试环境', level=2)
add_table(doc,
    ['项目', '要求'],
    [
        ['浏览器', 'Chrome 110+ 或 Edge 110+'],
        ['分辨率', '1366x768（最低）或 1920x1080（推荐）'],
        ['系统地址', '由项目负责人提供访问 URL'],
        ['外设', '耳机/麦克风（音频转写任务测试需要）'],
    ],
    col_widths=[4, 11])

add_heading(doc, '2.2  测试账号', level=2)
add_body(doc, '使用以下账号登录，按角色分层测试：')
add_table(doc,
    ['角色', '邮箱', '密码', '测试用途'],
    [
        ['学员 001', 'stu001@student.exam.local', 'abcd2345', '主测试账号（练习+考试）'],
        ['学员 002', 'stu002@student.exam.local', 'efgh6789', '辅助账号（多学员场景）'],
        ['教师', 'teacher01@exam.local', '随机生成', '教师端功能测试'],
        ['学校管理员', 'school@exam.local', '随机生成', '管理端功能测试'],
        ['超级管理员', 'admin@exam.local', '随机生成', '全部权限测试'],
        ['题库编辑', 'editor01@exam.local', 'Editor@2026', '题库编辑测试'],
        ['题库审核', 'reviewer01@exam.local', 'Review@2026', '题目审核测试'],
    ],
    col_widths=[2.5, 5, 3.5, 4])

add_heading(doc, '2.3  前置数据要求', level=2)
add_bullet(doc, '数据库已完成种子数据初始化（题库+任务模板+考试安排）')
add_bullet(doc, '至少有一场考试处于「待开始」或「进行中」状态')
add_bullet(doc, '练习题库中每种题型至少有 5 道题目')
add_bullet(doc, '实操任务模板已配置（含标注素材图片）')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 第三章  测试用例
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第三章  测试用例', level=1)

add_heading(doc, '3.1  登录认证模块', level=2)

add_body(doc, 'TC-AUTH-01：学员正常登录', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/login'],
        ['测试步骤', '1. 打开系统登录页\n2. 输入 stu001@student.exam.local / abcd2345\n3. 点击「登录」'],
        ['预期结果', '成功登录，自动跳转到 /student/home（学员首页）'],
        ['验证要点', '导航栏显示学员菜单（首页/练习/实操/错题本/考试/成绩）'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-AUTH-02：错误密码登录', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/login'],
        ['测试步骤', '1. 输入正确邮箱 + 错误密码\n2. 点击「登录」'],
        ['预期结果', '提示「邮箱或密码错误」，停留在登录页'],
        ['验证要点', '不跳转、不泄露是邮箱错还是密码错'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-AUTH-03：未登录访问受保护页面', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/home（未登录直接访问）'],
        ['测试步骤', '1. 清除浏览器 Cookie\n2. 直接访问 /student/home'],
        ['预期结果', '自动重定向到 /login'],
        ['验证要点', '不出现 404 或空白页，不泄露页面内容'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-AUTH-04：角色权限隔离', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/admin/dashboard（用学员账号登录后访问）'],
        ['测试步骤', '1. 用学员账号登录\n2. 手动输入 /admin/dashboard 访问'],
        ['预期结果', '提示无权限或重定向回学员首页'],
        ['验证要点', '学员不能访问管理端任何页面'],
    ],
    col_widths=[3, 12])

add_heading(doc, '3.2  学员端-理论练习模块', level=2)

add_body(doc, 'TC-PRACTICE-01：进入练习并答题', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/practice'],
        ['测试步骤', '1. 用 stu001 登录\n2. 点击「理论练习」\n3. 依次答题，答对和答错各一次'],
        ['预期结果', '答题后即时显示对错反馈；答对显示绿色「✓ 做对了！」；答错显示红色「✗ 答错了」并展示正确答案'],
        ['验证要点', '反馈文案温和鼓励、正确答案清晰可见'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-PRACTICE-02：错题自动收集', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/practice → /student/wrong'],
        ['测试步骤', '1. 在练习中故意答错 2 道题\n2. 进入「错题本」查看'],
        ['预期结果', '错题本中显示刚答错的题目'],
        ['验证要点', '错题可重做，做对后自动移出'],
    ],
    col_widths=[3, 12])

add_heading(doc, '3.3  学员端-实操任务模块', level=2)

add_body(doc, 'TC-TASK-01：图片清洗任务', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/task'],
        ['测试步骤', '1. 进入实操练习\n2. 选择「图片清洗」任务\n3. 从图片组中剔除错误样本（如把足球从篮球集合中删掉）\n4. 提交'],
        ['预期结果', '系统自动评分，显示得分和正确/错误清洗明细'],
        ['验证要点', '评分合理（正确剔除得分，误删扣分）'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-TASK-02：矩形框标注任务', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/task'],
        ['测试步骤', '1. 选择「方框标注」任务\n2. 在交通场景图上用矩形框标出指定目标（如人物、动物）\n3. 提交'],
        ['预期结果', '系统按 IoU 自动评分，标注越准分越高'],
        ['验证要点', '可以拖动调整框位置、删除标注后重新画'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-TASK-03：点标注任务', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/task'],
        ['测试步骤', '1. 选择「点标注」任务\n2. 在图片上点击指定目标位置（如红灯中心）\n3. 提交'],
        ['预期结果', '系统按距离精度评分'],
        ['验证要点', '点击位置可微调，多次点击只保留最后一个'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-TASK-04：文本情感标注任务', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/task'],
        ['测试步骤', '1. 选择「文本情感标注」任务\n2. 逐条判断文本的情感倾向（正面/中性/负面）\n3. 提交'],
        ['预期结果', '系统按标注准确率评分'],
        ['验证要点', '可以修改之前的标注，提交前可逐条检查'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-TASK-05：音频转写任务', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/task'],
        ['测试步骤', '1. 选择「音频转写」任务\n2. 点击播放按钮听音频\n3. 在输入框中输入听到的文字\n4. 提交'],
        ['预期结果', '系统按字符错误率(CER)评分'],
        ['验证要点', '音频可重复播放、输入框正常编辑'],
    ],
    col_widths=[3, 12])

doc.add_page_break()

add_heading(doc, '3.4  学员端-考试模块', level=2)

add_body(doc, 'TC-EXAM-01：开始考试', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/exams'],
        ['测试步骤', '1. 用 stu001 登录\n2. 点击「考试」\n3. 找到一场「进行中」的考试\n4. 点击「开始考试」'],
        ['预期结果', '进入考试答题界面，顶部显示倒计时'],
        ['验证要点', '考试界面克制严肃，不显示答题反馈'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-EXAM-02：考试中自动保存', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/exams（答题中）'],
        ['测试步骤', '1. 答到第 3 题\n2. 刷新浏览器页面\n3. 重新进入考试'],
        ['预期结果', '之前选的答案仍在，倒计时不重置'],
        ['验证要点', '进度恢复正确，倒计时连续'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-EXAM-03：手动交卷', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/exams（答题中）'],
        ['测试步骤', '1. 答完部分题目\n2. 点击「交卷」按钮\n3. 确认交卷'],
        ['预期结果', '成功交卷，跳转到成绩页或提示等待发布'],
        ['验证要点', '未答的题目不报错、按空值处理'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-EXAM-04：超时自动交卷', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/exams（答题中）'],
        ['测试步骤', '1. 开始考试后不操作\n2. 等待倒计时归零'],
        ['预期结果', '系统自动提交，考试状态变为「已过期」'],
        ['验证要点', '服务端时间锁生效，已完成答案正常评分'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-EXAM-05：未到考试时间不可开始', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/exams'],
        ['测试步骤', '1. 找一场「未开始」的考试\n2. 尝试点击开始'],
        ['预期结果', '提示考试尚未开始，或开始按钮不可点击'],
        ['验证要点', '服务端校验 exam_start_at <= NOW()'],
    ],
    col_widths=[3, 12])

add_heading(doc, '3.5  学员端-成绩查询模块', level=2)

add_body(doc, 'TC-RESULT-01：查看已发布成绩', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/student/results'],
        ['测试步骤', '1. 用 stu001 登录（需有已发布成绩）\n2. 点击「成绩」'],
        ['预期结果', '显示成绩列表（考试名称、总分、是否通过）'],
        ['验证要点', '未发布的成绩不显示'],
    ],
    col_widths=[3, 12])

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 3.6 - 3.7 教师端 & 管理端
# ═════════════════════════════════════════════════════════════

add_heading(doc, '3.6  教师端模块', level=2)

add_body(doc, 'TC-TEACHER-01：教师仪表盘', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/teacher/dashboard（或教师登录后首页）'],
        ['测试步骤', '1. 用 teacher01 登录\n2. 查看仪表盘'],
        ['预期结果', '显示所辖班级的学员人数、练习进度、考试情况'],
        ['验证要点', '只能看到自己班级的数据'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-TEACHER-02：教师查看学员列表', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/teacher/students'],
        ['测试步骤', '1. 进入学员列表\n2. 查看学员练习进度'],
        ['预期结果', '展示学员姓名、练习题数、正确率、任务完成数'],
        ['验证要点', '数据非空（前提：学员有练习记录）'],
    ],
    col_widths=[3, 12])

add_heading(doc, '3.7  管理端模块', level=2)

add_body(doc, 'TC-ADMIN-01：管理端工作台', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/admin/dashboard'],
        ['测试步骤', '1. 用 admin 登录\n2. 查看工作台概况'],
        ['预期结果', '显示系统统计（用户数、题目数、考试数等）'],
        ['验证要点', '左侧导航菜单完整展示'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-ADMIN-02：成绩复核-查看详情', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/admin/results'],
        ['测试步骤', '1. 进入成绩管理\n2. 找到一条成绩记录\n3. 点击「复核」查看详情'],
        ['预期结果', '显示总分 + 6 项分项分数 + 逐题答题对比'],
        ['验证要点', '学员答案、标准答案、得分、评分详情均可见'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-ADMIN-03：成绩复核-确认发布', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/admin/results（复核详情页）'],
        ['测试步骤', '1. 在复核详情页\n2. 填写复核说明\n3. 点击「确认并发布」'],
        ['预期结果', '成绩状态变为「已发布」，学员可查看成绩'],
        ['验证要点', '审计日志记录此操作'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-ADMIN-04：成绩复核-手动调整分数', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/admin/results（复核详情页）'],
        ['测试步骤', '1. 点击「调整分数」按钮\n2. 修改某一项分数\n3. 填写调整原因（至少 5 字）\n4. 点击「保存调整」'],
        ['预期结果', '总分自动重算，通过/不通过状态实时预判，保存成功'],
        ['验证要点', '变化项暖橙高亮、原始分删除线展示、审计日志含调整前后对比'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-ADMIN-05：成绩复核-调整校验', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/admin/results（复核详情页）'],
        ['测试步骤', '1. 点击「调整分数」\n2. 不修改任何值，直接保存\n3. 修改值但不填原因\n4. 修改值但总分超过满分'],
        ['预期结果', '场景2：提示「至少修改一项」；场景3：提示「原因至少5字」；场景4：红色警告'],
        ['验证要点', '校验完整，不允许异常数据提交'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-ADMIN-06：题库导入', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/admin/import'],
        ['测试步骤', '1. 进入题库导入页面\n2. 上传 Word 题库文件\n3. 确认导入'],
        ['预期结果', '题目成功导入，显示导入数量和失败项'],
        ['验证要点', '导入的题目在练习题库中可查看'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-ADMIN-07：考务安排', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/admin/exam-schedules'],
        ['测试步骤', '1. 进入考务安排\n2. 创建一场考试（选择班级、试卷、时间）\n3. 保存'],
        ['预期结果', '考试安排成功创建，学员端考试列表可见'],
        ['验证要点', '考试时间范围正确设置'],
    ],
    col_widths=[3, 12])

add_body(doc, 'TC-ADMIN-08：审计日志查询', bold=True)
add_table(doc,
    ['项目', '内容'],
    [
        ['测试 URL', '/admin/audit'],
        ['测试步骤', '1. 进入审计日志\n2. 按时间/操作类型筛选\n3. 查看日志详情'],
        ['预期结果', '显示操作人、时间、操作类型、详情'],
        ['验证要点', '成绩复核/调整操作有完整记录'],
    ],
    col_widths=[3, 12])

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 第四章  缺陷记录规范
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第四章  缺陷记录规范', level=1)

add_heading(doc, '4.1  缺陷严重等级', level=2)
add_table(doc,
    ['等级', '标识', '定义', '示例'],
    [
        ['致命', 'S1', '系统崩溃/数据丢失/核心功能完全不可用', '登录失败、考试无法开始、提交报500'],
        ['严重', 'S2', '核心功能异常但可绕过', '评分错误、部分题型不能提交'],
        ['一般', 'S3', '非核心功能异常或交互体验问题', '错题本未更新、表格排序错误'],
        ['轻微', 'S4', 'UI显示瑕疵/文案错误', '错别字、间距不一致、图标错位'],
    ],
    col_widths=[2, 1.5, 5, 6.5])

add_heading(doc, '4.2  缺陷记录模板', level=2)
add_body(doc, '每发现一个 Bug，请按以下格式记录（详见测试记录文档）：')
add_table(doc,
    ['字段', '说明', '示例'],
    [
        ['Bug ID', '唯一编号', 'BUG-2026-001'],
        ['发现日期', 'YYYY-MM-DD', '2026-07-27'],
        ['发现人', '测试人员姓名', '张三'],
        ['角色/账号', '用什么账号测试时发现', '学员 stu001'],
        ['页面 URL', '出问题的页面地址', '/student/exams'],
        ['严重等级', 'S1/S2/S3/S4', 'S2'],
        ['Bug 标题', '一句话描述', '交卷后页面空白'],
        ['问题描述', '详细描述（步骤、实际结果、预期结果）', '答完10题点交卷后白屏，控制台报错 TypeError'],
        ['截图', '粘贴问题截图', '（在此处插入截图）'],
        ['复现步骤', '编号步骤', '1. 登录stu001 2. 开始考试 3. 答题 4. 点交卷'],
        ['状态', '新建/确认/修复/关闭', '新建'],
    ],
    col_widths=[3, 4, 8])

add_heading(doc, '4.3  截图要求', level=2)
add_bullet(doc, '必须截取完整浏览器窗口（含地址栏 URL）')
add_bullet(doc, '如有控制台错误，需附加 F12 开发者工具截图')
add_bullet(doc, '截图文件命名：BugID_序号.png（如 BUG-2026-001_1.png）')
add_bullet(doc, '截图粘贴到测试记录文档对应的 Bug 记录处')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 第五章  测试执行计划
# ═════════════════════════════════════════════════════════════

add_heading(doc, '第五章  测试执行计划', level=1)

add_heading(doc, '5.1  测试阶段划分', level=2)
add_table(doc,
    ['阶段', '范围', '通过标准'],
    [
        ['第一轮：冒烟测试', '登录 + 各角色首页可访问', '所有 P0 页面不报500'],
        ['第二轮：功能测试', '本文档全部测试用例', 'P0/S1 全部通过，S2 缺陷已记录'],
        ['第三轮：回归测试', '修复后的缺陷验证', '已修复 Bug 不复现'],
        ['第四轮：验收测试', '全链路走通', '无 S1/S2 未关闭缺陷'],
    ],
    col_widths=[3.5, 6, 5.5])

add_heading(doc, '5.2  测试报告输出', level=2)
add_body(doc, '测试完成后，请输出以下内容：')
add_bullet(doc, '填写完整的「测试记录文档」（每个用例标记 通过/失败/阻塞）')
add_bullet(doc, '汇总缺陷列表（按严重等级排序）')
add_bullet(doc, '测试覆盖率统计：已测用例数 / 总用例数')
add_bullet(doc, '遗留问题与风险评估')

# ── save ─────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Test guide saved to: {OUTPUT}")
print(f"File size: {os.path.getsize(OUTPUT) / 1024:.0f} KB")
